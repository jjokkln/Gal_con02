"""
Export functionality for PDF and DOCX generation
"""

import io
from typing import Dict, Any
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
class ProfileExporter:
    def __init__(self):
        pass
    
    def prepare_for_print(self, html_string: str) -> str:
        """
        Prepare HTML for browser-based PDF printing with optimized CSS
        """
        # Enhanced print CSS for better page breaks and layout
        enhanced_print_css = """
        <style>
            @media print {
                * {
                    -webkit-print-color-adjust: exact !important;
                    print-color-adjust: exact !important;
                    color-adjust: exact !important;
                }
                @page {
                    size: A4;
                    margin: 0;
                }
                body {
                    margin: 0;
                    padding: 0;
                }
            }
            @media screen {
                body {
                    background: #525252;
                    padding: 20px;
                }
                .profile-container, .container {
                    box-shadow: 0 0 20px rgba(0,0,0,0.3);
                    margin: 0 auto;
                    background: white;
                }
            }
        </style>
        <script>
            // Auto-trigger print dialog on page load
            window.onload = function() {
                // Give browser time to render
                setTimeout(function() {
                    window.print();
                }, 500);
            };
        </script>
        """
        
        # Insert enhanced CSS and script before closing head tag
        if "</head>" in html_string:
            html_with_enhancements = html_string.replace("</head>", f"{enhanced_print_css}</head>")
        else:
            html_with_enhancements = enhanced_print_css + html_string
            
        return html_with_enhancements
    
    def html_to_pdf(self, html_string: str) -> str:
        """
        Return HTML string for browser-based PDF printing
        This method prepares HTML with print-friendly CSS
        """
        # Add print-friendly CSS to the HTML
        print_css = """
        <style>
            @media print {
                @page {
                    size: A4;
                    margin: 2cm;
                }
                body {
                    -webkit-print-color-adjust: exact;
                    print-color-adjust: exact;
                }
            }
        </style>
        """
        
        # Insert CSS before closing head tag
        if "</head>" in html_string:
            html_with_print_css = html_string.replace("</head>", f"{print_css}</head>")
        else:
            html_with_print_css = print_css + html_string
            
        return html_with_print_css
    
    def generate_pdf_from_data(self, data: Dict[str, Any], company: str = "galdora", template: str = "modern") -> bytes:
        """
        Generate PDF directly from CV data using ReportLab
        Supports multiple templates: "modern" and "classic"
        """
        try:
            pdf_buffer = io.BytesIO()
            doc = SimpleDocTemplate(pdf_buffer, pagesize=A4,
                                  topMargin=2*cm, bottomMargin=2*cm,
                                  leftMargin=2*cm, rightMargin=2*cm)
            
            # Company colors and logos
            company_colors_map = {
                "galdora": colors.HexColor("#1e3a8a"),
                "bejob": colors.HexColor("#059669")
            }
            primary_color = company_colors_map.get(company.lower(), company_colors_map["galdora"])
            
            # Company logo paths
            company_logos = {
                "galdora": "ressources/galdora.svg",
                "bejob": "ressources/bejob-logo.png"
            }
            logo_path = company_logos.get(company.lower())
            
            # Styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=primary_color,
                spaceAfter=12,
                alignment=TA_CENTER
            )
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=primary_color,
                spaceBefore=12,
                spaceAfter=6
            )
            normal_style = styles['Normal']
            bullet_style = ParagraphStyle(
                'Bullet',
                parent=normal_style,
                leftIndent=0.5*cm,
                bulletIndent=0.2*cm,
                spaceBefore=0.1*cm,
                bulletText='•'
            )
            
            # Build PDF content
            content = []

            # Header: Logo rechts zentriert, Name/Position links zentriert
            personal = data.get("personal", {})
            header_elements = []

            # Name und Position links zentriert (untereinander)
            left_content = []
            if personal.get("name"):
                left_content.append(Paragraph(personal["name"], title_style))
            if personal.get("position"):
                left_content.append(Paragraph(personal["position"], normal_style))

            if left_content:
                left_cell = Table([[left_content[0]]] + [[left_content[i]] for i in range(1, len(left_content))],
                                colWidths=[8*cm])
                left_cell.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                header_elements.append(left_cell)
            else:
                header_elements.append(Paragraph("", normal_style))

            # Logo rechts zentriert
            if logo_path and os.path.exists(logo_path):
                try:
                    if logo_path.lower().endswith('.svg'):
                        logo = Image(logo_path, width=4*cm, height=1.5*cm)
                    else:
                        logo = Image(logo_path, width=4*cm, height=1.5*cm)
                    header_elements.append(logo)
                except Exception as e:
                    header_elements.append(Paragraph("", normal_style))
            else:
                header_elements.append(Paragraph("", normal_style))

            # Header Table (Name/Position links, Logo rechts) - kürzer
            if len(header_elements) == 2:
                header_table = Table([header_elements], colWidths=[8*cm, 7*cm])
                header_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (0, 0), 'CENTER'),
                    ('ALIGN', (1, 0), (1, 0), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                content.append(header_table)
                content.append(Spacer(1, 0.5*cm))

            # Summary direkt unter Header zentriert (optional)
            show_summary = data.get("_export_options", {}).get("show_summary", True)
            if show_summary and personal.get("summary"):
                summary_style = ParagraphStyle(
                    'Summary',
                    parent=normal_style,
                    alignment=TA_CENTER,
                    fontSize=11,
                    spaceBefore=0.3*cm,
                    spaceAfter=0.5*cm,
                    textColor=colors.grey
                )
                content.append(Paragraph(personal["summary"], summary_style))
            
            # Experience
            experience = data.get("experience", [])
            limit_tasks = data.get("_export_options", {}).get("limit_tasks", True)
            
            if experience:
                content.append(Paragraph("Berufserfahrung", heading_style))
                for exp in experience:
                    exp_title = f"<b>{exp.get('position', '')}</b> bei {exp.get('company', '')}"
                    content.append(Paragraph(exp_title, normal_style))
                    
                    dates = f"{exp.get('start_date', '')} - {exp.get('end_date', '')}"
                    date_style = ParagraphStyle('Dates', parent=normal_style, textColor=colors.grey)
                    content.append(Paragraph(dates, date_style))
                    
                    # Tasks als Stichpunkte (max 5 wenn limit_tasks=True)
                    tasks = exp.get("tasks", [])
                    if tasks:
                        display_tasks = tasks[:5] if limit_tasks else tasks
                        for task in display_tasks:
                            bullet_para = Paragraph(f"• {task}", normal_style)
                            content.append(bullet_para)
                    elif exp.get("description"):
                        # Fallback für altes Format
                        content.append(Paragraph(exp["description"], normal_style))
                    content.append(Spacer(1, 0.3*cm))
            
            # Education
            education = data.get("education", [])
            if education:
                content.append(Paragraph("Ausbildung & Weiterbildung", heading_style))
                for edu in education:
                    edu_title = f"<b>{edu.get('degree', '')}</b> - {edu.get('institution', '')}"
                    content.append(Paragraph(edu_title, normal_style))
                    
                    dates = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}"
                    date_style = ParagraphStyle('Dates', parent=normal_style, textColor=colors.grey)
                    content.append(Paragraph(dates, date_style))
                    
                    if edu.get("description"):
                        content.append(Paragraph(edu["description"], normal_style))
                    content.append(Spacer(1, 0.3*cm))
            
            # Skills
            skills = data.get("skills", [])
            if skills:
                content.append(Paragraph("Fähigkeiten & Kompetenzen", heading_style))
                skills_text = ", ".join(skills)
                content.append(Paragraph(skills_text, normal_style))
                content.append(Spacer(1, 0.3*cm))
            
            # Languages
            languages = data.get("languages", [])
            if languages:
                content.append(Paragraph("Sprachen", heading_style))
                for lang in languages:
                    lang_text = f"<b>{lang.get('name', '')}</b> - {lang.get('level', '')}"
                    content.append(Paragraph(lang_text, normal_style))
                    content.append(Spacer(1, 0.2*cm))
                content.append(Spacer(1, 0.3*cm))
            
            # Build PDF
            doc.build(content)
            
            # Get PDF bytes
            pdf_bytes = pdf_buffer.getvalue()
            pdf_buffer.close()
            
            return pdf_bytes
            
        except Exception as e:
            raise Exception(f"Error generating PDF: {str(e)}")
    
    def generate_docx(self, data: Dict[str, Any], company: str = "galdora") -> bytes:
        """
        Generate DOCX file from CV data
        """
        try:
            doc = Document()
            
            # Set up document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            # Add header with name
            personal = data.get("personal", {})
            if personal.get("name"):
                header = doc.add_heading(personal["name"], 0)
                header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add contact information
            if any(personal.get(key) for key in ["email", "phone", "address", "linkedin"]):
                contact_para = doc.add_paragraph()
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                contact_info = []
                if personal.get("email"):
                    contact_info.append(f"📧 {personal['email']}")
                if personal.get("phone"):
                    contact_info.append(f"📞 {personal['phone']}")
                if personal.get("address"):
                    contact_info.append(f"📍 {personal['address']}")
                if personal.get("linkedin"):
                    contact_info.append(f"💼 {personal['linkedin']}")
                
                contact_para.add_run(" | ".join(contact_info))
            
            # Add summary
            if personal.get("summary"):
                doc.add_heading("Profil", level=1)
                summary_para = doc.add_paragraph(personal["summary"])
                summary_para.italic = True
            
            # Add experience
            experience = data.get("experience", [])
            limit_tasks = data.get("_export_options", {}).get("limit_tasks", True)
            
            if experience:
                doc.add_heading("Berufserfahrung", level=1)
                for exp in experience:
                    # Position and company
                    exp_heading = doc.add_heading(f"{exp.get('position', '')} bei {exp.get('company', '')}", level=2)
                    
                    # Dates
                    if exp.get("start_date") or exp.get("end_date"):
                        date_para = doc.add_paragraph()
                        date_para.add_run(f"{exp.get('start_date', '')} - {exp.get('end_date', '')}").italic = True
                    
                    # Tasks als Stichpunkte oder Description als Fallback
                    tasks = exp.get("tasks", [])
                    if tasks:
                        display_tasks = tasks[:5] if limit_tasks else tasks
                        for task in display_tasks:
                            doc.add_paragraph(task, style='List Bullet')
                    elif exp.get("description"):
                        desc_para = doc.add_paragraph(exp["description"])
            
            # Add education
            education = data.get("education", [])
            if education:
                doc.add_heading("Ausbildung & Weiterbildung", level=1)
                for edu in education:
                    # Degree and institution
                    edu_heading = doc.add_heading(f"{edu.get('degree', '')} - {edu.get('institution', '')}", level=2)
                    
                    # Dates
                    if edu.get("start_date") or edu.get("end_date"):
                        date_para = doc.add_paragraph()
                        date_para.add_run(f"{edu.get('start_date', '')} - {edu.get('end_date', '')}").italic = True
                    
                    # Description
                    if edu.get("description"):
                        desc_para = doc.add_paragraph(edu["description"])
            
            # Add skills
            skills = data.get("skills", [])
            if skills:
                doc.add_heading("Fähigkeiten & Kompetenzen", level=1)
                skills_para = doc.add_paragraph()
                skills_para.add_run(", ".join(skills))
            
            # Add languages
            languages = data.get("languages", [])
            if languages:
                doc.add_heading("Sprachen", level=1)
                for lang in languages:
                    lang_para = doc.add_paragraph()
                    lang_para.add_run(f"{lang.get('name', '')} - {lang.get('level', '')}").bold = True
            
            # Add certifications
            certifications = data.get("certifications", [])
            if certifications:
                doc.add_heading("Zertifizierungen", level=1)
                for cert in certifications:
                    cert_para = doc.add_paragraph()
                    cert_para.add_run(f"{cert.get('name', '')} - {cert.get('issuer', '')}").bold = True
                    if cert.get("date"):
                        cert_para.add_run(f" ({cert['date']})")
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            return doc_bytes.getvalue()
            
        except Exception as e:
            raise Exception(f"Error generating DOCX: {str(e)}")


# Convenience functions for direct usage
def html_to_pdf(html_string: str) -> bytes:
    """Convert HTML to PDF"""
    exporter = ProfileExporter()
    return exporter.html_to_pdf(html_string)

def generate_docx(data: Dict[str, Any], company: str = "galdora") -> bytes:
    """Generate DOCX from CV data"""
    exporter = ProfileExporter()
    return exporter.generate_docx(data, company)
